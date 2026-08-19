"""
AGROVIAL MRA - Presupuestos Web v2.0 (Render Deploy)
"""
import os, io, json, re
from datetime import datetime
from flask import Flask, render_template, request, jsonify, send_file, send_from_directory
from docx import Document

try:
    from fpdf import FPDF
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

app = Flask(__name__)
APP_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE = os.path.join(APP_DIR, 'PRESUPUESTO WORD.docx')
LOGO_FILE = os.path.join(APP_DIR, 'static', 'logo.png')

_counter = 951
_saved = {}


def parse_money(value):
    s = str(value).strip().replace('Gs.', '').replace(' ', '')
    if not s:
        return 0
    s = s.replace('.', '').replace(',', '')
    try:
        return int(s)
    except ValueError:
        return 0


def money(n):
    return f'{int(round(n)):,}'.replace(',', '.')


def next_number():
    global _counter
    n = _counter
    _counter += 1
    return n


@app.route('/')
def index():
    return render_template('index.html', has_pdf=HAS_PDF, num=f'{_counter:07d}',
                           date=datetime.now().strftime('%d/%m/%Y'))


@app.route('/logo')
def logo():
    if os.path.exists(LOGO_FILE):
        return send_from_directory(os.path.join(APP_DIR, 'static'), 'logo.png')
    return '', 404


@app.route('/api/next_num')
def next_num():
    return jsonify({'num': f'{_counter:07d}'})


@app.route('/api/generate_word', methods=['POST'])
def generate_word():
    if not os.path.exists(TEMPLATE):
        return jsonify({'error': 'Falta PRESUPUESTO WORD.docx'}), 400
    data = request.json
    items = [i for i in data.get('items', []) if i.get('name', '').strip()]
    if not items:
        return jsonify({'error': 'Sin items'}), 400

    grand = sum(float(i.get('qty', 0) or 0) * parse_money(i.get('price', 0)) for i in items)
    num = next_number()
    date = data.get('date', datetime.now().strftime('%d/%m/%Y'))
    client = data.get('client', '')

    doc = Document(TEMPLATE)
    for p in doc.paragraphs:
        if 'Nombre o Razón Social:' in p.text:
            p.text = f'Nombre o Razon Social:          {client} '
        if re.search(r'Fecha:', p.text) and 'Presupuesto' in p.text:
            p.text = f'Fecha: {date} Presupuesto {num:07d}'
        elif p.text.strip().startswith('Total:'):
            p.text = f'Total: {money(grand)}'

    if doc.tables:
        t = doc.tables[0]
        for r in range(1, min(13, len(t.rows))):
            item = items[r - 1] if r - 1 < len(items) else {}
            name = item.get('name', '')
            q = float(item.get('qty', 0) or 0)
            p_val = parse_money(item.get('price', 0))
            tot = q * p_val
            cells = t.rows[r].cells
            cells[0].text = str(name)
            cells[1].text = money(q) if name else ''
            cells[2].text = money(p_val) if name else ''
            cells[3].text = money(tot) if name else ''

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        as_attachment=True,
        download_name=f'PRESUPUESTO {num:07d}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/api/generate_pdf', methods=['POST'])
def generate_pdf_route():
    if not HAS_PDF:
        return jsonify({'error': 'Instale fpdf2'}), 400
    data = request.json
    items = [i for i in data.get('items', []) if i.get('name', '').strip()]
    if not items:
        return jsonify({'error': 'Sin items'}), 400

    grand = sum(float(i.get('qty', 0) or 0) * parse_money(i.get('price', 0)) for i in items)
    num = next_number()
    date = data.get('date', datetime.now().strftime('%d/%m/%Y'))
    client = data.get('client', '')
    orientation = data.get('orientation', 'L')

    is_landscape = orientation == 'L'
    pdf = FPDF(orientation=orientation, unit='mm', format='A4')
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    W = 297 if is_landscape else 210

    pdf.set_fill_color(13, 27, 42)
    pdf.rect(0, 0, W, 35, 'F')

    x_start = 10
    if os.path.exists(LOGO_FILE):
        try:
            pdf.image(LOGO_FILE, x=10, y=5, h=25)
            x_start = 42
        except Exception:
            pass

    pdf.set_font('Helvetica', 'B', 16)
    pdf.set_text_color(255, 255, 255)
    pdf.set_xy(x_start, 8)
    pdf.cell(100, 8, 'AGROVIAL MRA', align='L')

    pdf.set_font('Helvetica', 'B', 14)
    pdf.set_text_color(136, 150, 166)
    pdf.set_xy(W - 97, 8)
    pdf.cell(87, 6, f'Presupuesto N\u00ba {num:07d}', align='R')

    pdf.set_fill_color(136, 150, 166)
    pdf.rect(0, 35, W, 2, 'F')

    pdf.set_y(42)
    pdf.set_font('Helvetica', '', 10)
    pdf.set_text_color(26, 26, 46)
    pdf.cell(40, 7, 'Cliente:')
    pdf.set_font('Helvetica', 'B', 10)
    pdf.cell(0, 7, client, ln=True)
    pdf.set_font('Helvetica', '', 10)
    pdf.cell(40, 7, 'Fecha:')
    pdf.cell(0, 7, date, ln=True)
    pdf.ln(3)

    if is_landscape:
        col_widths = [135, 30, 60, 60]
        headers = ['ARTICULO / SERVICIO', 'CANT.', 'PRECIO UNIT. (Gs.)', 'TOTAL (Gs.)']
        hdr_font, row_font, max_name = 9, 10, 50
    else:
        col_widths = [85, 20, 35, 35]
        headers = ['ARTICULO', 'CANT.', 'PRECIO UNIT.', 'TOTAL']
        hdr_font, row_font, max_name = 8, 8, 32

    pdf.set_fill_color(13, 27, 42)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font('Helvetica', 'B', hdr_font)
    for h, w in zip(headers, col_widths):
        pdf.cell(w, 7, h, border=0, fill=True, align='C')
    pdf.ln()

    for idx, item in enumerate(items):
        q = float(item.get('qty', 0) or 0)
        p_val = parse_money(item.get('price', 0))
        tot = q * p_val
        name = item.get('name', '')
        if idx % 2 == 0:
            pdf.set_fill_color(244, 247, 251)
        else:
            pdf.set_fill_color(255, 255, 255)
        pdf.set_text_color(26, 26, 46)
        pdf.set_font('Helvetica', '', row_font)
        pdf.cell(col_widths[0], 7, name[:max_name], border=0, fill=True, align='L')
        pdf.set_font('Helvetica', 'B', row_font)
        pdf.cell(col_widths[1], 7, money(q), border=0, fill=True, align='C')
        pdf.cell(col_widths[2], 7, money(p_val), border=0, fill=True, align='R')
        pdf.cell(col_widths[3], 7, money(tot), border=0, fill=True, align='R')
        pdf.ln()

    pdf.ln(5)
    total_w = sum(col_widths)
    pdf.set_draw_color(197, 205, 216)
    pdf.line(W - total_w - 10, pdf.get_y(), W - 10, pdf.get_y())
    pdf.ln(5)

    pdf.set_fill_color(13, 27, 42)
    pdf.set_x(W - total_w - 10)
    pdf.set_font('Helvetica', 'B', 11)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(total_w, 10, f'TOTAL:  Gs. {money(grand)}', fill=True, align='R')

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)

    return send_file(
        buf,
        as_attachment=True,
        download_name=f'PRESUPUESTO {num:07d}.pdf',
        mimetype='application/pdf'
    )


@app.route('/api/save', methods=['POST'])
def save_presupuesto():
    data = request.json
    num = data.get('num', '0000000')
    _saved[num] = data
    return jsonify({'success': True, 'num': num})


@app.route('/api/saved_list')
def saved_list():
    result = []
    for num, d in _saved.items():
        result.append({'file': num, 'client': d.get('client', ''), 'date': d.get('date', ''), 'num': d.get('num', '')})
    return jsonify(result)


@app.route('/api/load/<filename>')
def load_presupuesto(filename):
    if filename not in _saved:
        return jsonify({'error': 'No encontrado'}), 404
    return jsonify(_saved[filename])


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(debug=False, host='0.0.0.0', port=port)
